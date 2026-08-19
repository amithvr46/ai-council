"""Text extraction from uploaded source material.

Deliberately narrow (contract §2C): PDF, DOCX, TXT, MD and code files. Unusual
formats are refused with a clear message rather than half-supported — a file
that cannot be parsed must never become an empty document that looks like an
empty resume.
"""

import io
from dataclasses import dataclass

MAX_BYTES = 10 * 1024 * 1024  # 10 MB
MAX_CHARS = 120_000  # ~30k tokens; a resume is ~4k chars, a JD ~6k

TEXT_EXTENSIONS = {
    ".txt", ".md", ".markdown", ".rst",
    ".py", ".js", ".ts", ".tsx", ".jsx", ".go", ".rs", ".java", ".rb", ".sh",
    ".ps1", ".yaml", ".yml", ".json", ".toml", ".ini", ".cfg", ".conf",
    ".tf", ".tfvars", ".hcl", ".sql", ".xml", ".html", ".css", ".env",
    ".dockerfile", ".gitignore",
}


class ExtractionError(Exception):
    """Unparseable input. Surfaces to the user; never silently empty."""


@dataclass
class Extracted:
    text: str
    char_count: int
    truncated: bool
    detected_kind: str  # pdf | docx | text


def _suffix(filename: str) -> str:
    name = filename.lower()
    if "." not in name:
        return ""
    return name[name.rindex(".") :]


def extract(filename: str, data: bytes) -> Extracted:
    if len(data) > MAX_BYTES:
        raise ExtractionError(
            f"{filename} is {len(data) // 1024 // 1024} MB; the limit is "
            f"{MAX_BYTES // 1024 // 1024} MB"
        )
    if not data:
        raise ExtractionError(f"{filename} is empty")

    suffix = _suffix(filename)
    if suffix == ".pdf":
        text, kind = _extract_pdf(filename, data), "pdf"
    elif suffix in (".docx", ".doc"):
        text, kind = _extract_docx(filename, data), "docx"
    elif suffix in TEXT_EXTENSIONS or suffix == "":
        text, kind = _extract_text(filename, data), "text"
    else:
        raise ExtractionError(
            f"{filename}: unsupported file type '{suffix}'. Supported: PDF, DOCX, "
            f"TXT, MD and common code/config files."
        )

    text = text.strip()
    if not text:
        raise ExtractionError(
            f"{filename}: no text could be extracted. If this is a scanned or "
            f"image-only PDF, it needs OCR, which is not supported."
        )

    truncated = len(text) > MAX_CHARS
    if truncated:
        text = text[:MAX_CHARS]
    return Extracted(text=text, char_count=len(text), truncated=truncated, detected_kind=kind)


def _extract_pdf(filename: str, data: bytes) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as e:  # pragma: no cover - dependency is declared
        raise ExtractionError("PDF support requires pypdf") from e
    try:
        reader = PdfReader(io.BytesIO(data))
        return "\n\n".join((page.extract_text() or "") for page in reader.pages)
    except Exception as e:
        raise ExtractionError(f"{filename}: could not read PDF ({type(e).__name__})") from e


def _extract_docx(filename: str, data: bytes) -> str:
    try:
        import docx
    except ImportError as e:  # pragma: no cover - dependency is declared
        raise ExtractionError("DOCX support requires python-docx") from e
    try:
        document = docx.Document(io.BytesIO(data))
    except Exception as e:
        raise ExtractionError(f"{filename}: could not read DOCX ({type(e).__name__})") from e

    parts = [p.text for p in document.paragraphs]
    # Resumes frequently use tables for layout; losing them loses the resume.
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def _extract_text(filename: str, data: bytes) -> str:
    for encoding in ("utf-8", "utf-16", "latin-1"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    raise ExtractionError(f"{filename}: could not decode as text")
