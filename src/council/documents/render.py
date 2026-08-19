"""Submission-ready DOCX output (contract Amendment A13).

python-docx rather than a Node renderer: this ships as a single Python service
the user runs on their own machine, and a second runtime for one output format
is a dependency they would have to install and keep working. python-docx is
already here for extraction, and a resume needs headings, bullets and a contact
line — not a layout engine.

The formatting target is a resume that survives an ATS parse: real paragraph
styles, no text boxes, no tables used for layout, no columns, and nothing
load-bearing in headers or footers. Everything an ATS needs is in the document
body, in reading order.
"""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

from council.documents.schemas import ResumeDraft

_INK = RGBColor(0x11, 0x18, 0x27)
_MUTED = RGBColor(0x4B, 0x55, 0x63)


def _configure(document: Document) -> None:
    for section in document.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.65)
        section.right_margin = Inches(0.65)

    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10)
    normal.paragraph_format.space_after = Pt(2)
    normal.paragraph_format.line_spacing = 1.0


def _section_heading(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(10)
    paragraph.paragraph_format.space_after = Pt(3)
    run = paragraph.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(10.5)
    run.font.color.rgb = _INK
    # A bottom border, not a table — a table here would break ATS parsing.
    borders = paragraph.paragraph_format.element.get_or_add_pPr()
    from docx.oxml.ns import qn
    from docx.oxml.parser import OxmlElement

    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "9CA3AF")
    pbdr.append(bottom)
    borders.append(pbdr)


def _bullet(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(text, style="List Bullet")
    paragraph.paragraph_format.space_after = Pt(2)
    # Tight hanging indent. The built-in style leaves a wide gap between the
    # glyph and the text, which wastes horizontal space on a document whose
    # whole constraint is space.
    paragraph.paragraph_format.left_indent = Inches(0.21)
    paragraph.paragraph_format.first_line_indent = Inches(-0.14)


def render_docx(
    draft: ResumeDraft,
    path: str | Path,
    *,
    name: str = "",
    contact: str = "",
) -> Path:
    """Write the draft to `path` and return it.

    `name` and `contact` are passed in rather than generated: they are personal
    details that belong to the user's own records, and a model has no business
    producing an email address or phone number.
    """
    path = Path(path)
    path.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    _configure(document)

    if name:
        heading = document.add_paragraph()
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        heading.paragraph_format.space_after = Pt(0)
        run = heading.add_run(name.upper())
        run.bold = True
        run.font.size = Pt(17)
        run.font.color.rgb = _INK

    if draft.headline:
        line = document.add_paragraph()
        line.alignment = WD_ALIGN_PARAGRAPH.CENTER
        line.paragraph_format.space_after = Pt(1)
        run = line.add_run(draft.headline)
        run.font.size = Pt(10.5)
        run.font.color.rgb = _MUTED

    if contact:
        line = document.add_paragraph()
        line.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = line.add_run(contact)
        run.font.size = Pt(9)
        run.font.color.rgb = _MUTED

    if draft.summary:
        _section_heading(document, "Professional Summary")
        document.add_paragraph(draft.summary)

    if draft.skills:
        _section_heading(document, "Technical Skills")
        for group, terms in draft.skills.items():
            if not terms:
                continue
            paragraph = document.add_paragraph()
            label = paragraph.add_run(f"{group}: ")
            label.bold = True
            paragraph.add_run(", ".join(terms))

    if draft.roles:
        _section_heading(document, "Professional Experience")
        for role in draft.roles:
            header = document.add_paragraph()
            header.paragraph_format.space_before = Pt(6)
            header.paragraph_format.space_after = Pt(1)
            title = header.add_run(f"{role.title} | {role.employer}")
            title.bold = True
            title.font.size = Pt(10.5)
            tail = " | ".join(p for p in (role.location, role.dates) if p)
            if tail:
                meta = header.add_run(f" | {tail}")
                meta.font.size = Pt(9.5)
                meta.font.color.rgb = _MUTED
            for bullet in role.bullets:
                _bullet(document, bullet)

    if draft.projects:
        _section_heading(document, "Technical Projects")
        for project in draft.projects:
            header = document.add_paragraph()
            header.paragraph_format.space_before = Pt(6)
            header.paragraph_format.space_after = Pt(1)
            run = header.add_run(project.name)
            run.bold = True
            if project.context:
                meta = header.add_run(f" | {project.context}")
                meta.font.size = Pt(9.5)
                meta.font.color.rgb = _MUTED
            for bullet in project.bullets:
                _bullet(document, bullet)

    if draft.education:
        _section_heading(document, "Education")
        for line in draft.education:
            document.add_paragraph(line)

    document.save(path)
    return path
