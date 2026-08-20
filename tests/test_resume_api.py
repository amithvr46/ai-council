"""The One-Step resume path through the API.

The product promise being tested: career sources + a JD + one line of plain
English in, a downloadable DOCX out. No mode, no model choice, no bullet
approval, no terminal.

The load-bearing distinction is in §3 of the contract — one instruction carries
a DURABLE career fact and a REQUEST-ONLY preference, and confusing them fails
in both directions.
"""

import io

from docx import Document as DocxDocument
from fastapi.testclient import TestClient
from sqlalchemy import select

import council.api.main as api
from council import outcomes
from council.api.main import app
from council.db.models import ArtifactRow, DocumentRow
from council.db.session import session_scope
from council.documents.instructions import parse
from council.documents.profile import AUTHORITY_USER_STATEMENT
from council.documents.schemas import ExperienceSelection, ResumeDraft, ResumeRole, RoleEmphasis
from council.documents.workflow import ResumeWorkflow
from council.engine.prompts import default_registry
from council.engine.schemas import CombinedCheck, Synthesis
from tests.fakes import FakeProvider

GCP_JD = (
    "Infrastructure Engineer. Primarily Google Cloud Platform: GKE, Cloud Run and "
    "Cloud SQL. Write Terraform for all infrastructure, build CI/CD pipelines and "
    "support production Kubernetes workloads."
)


def _selection():
    return ExperienceSelection(
        target_summary="Cloud infrastructure and delivery",
        priority_themes=["terraform", "kubernetes"],
        roles=[RoleEmphasis(employer="Acme", title="Cloud Engineer", bullet_budget=3)],
    )


def _draft():
    return ResumeDraft(
        headline="Cloud / DevOps Engineer",
        summary="Cloud engineer working across Azure infrastructure and delivery pipelines.",
        skills={"Cloud": ["Azure", "Terraform", "Kubernetes"]},
        roles=[
            ResumeRole(
                title="Cloud Engineer",
                employer="Acme",
                bullets=[
                    "Reviewed Terraform plans ahead of Azure infrastructure changes and "
                    "reconciled drift back to the approved configuration.",
                    "Supported AKS workloads during releases by checking pod health, "
                    "restart patterns and rollout status.",
                ],
            )
        ],
    )


def _fake_workflow():
    drafter = FakeProvider("drafter", [_selection(), _draft()])
    reviewer = FakeProvider("reviewer", [None])
    return ResumeWorkflow(
        {"drafter": drafter, "reviewer": reviewer},
        default_registry(),
        draft_provider="drafter",
        review_provider="reviewer",
        flagship_models={"drafter": "fake", "reviewer": "fake"},
        cheap_models={"drafter": "fake-cheap", "reviewer": "fake-cheap"},
    )


def _client(monkeypatch) -> TestClient:
    """Swap in fakes so no provider is ever contacted."""
    import council.engine.factory as factory

    engine = None  # the ask pipeline is not exercised here
    monkeypatch.setattr(factory, "build_resume_workflow", _fake_workflow)
    monkeypatch.setattr(
        "council.engine.factory.build_engine",
        lambda *a, **k: _ask_engine(),
    )
    client = TestClient(app, raise_server_exceptions=True)
    client.__enter__()
    monkeypatch.setattr(api, "engine", engine)
    return client


def _ask_engine():
    from council.engine.pipeline import CouncilEngine

    return CouncilEngine(
        {
            "openai": FakeProvider(
                "openai",
                ["a", CombinedCheck(agreement="agree", disagreement_type="none", summary="s"),
                 Synthesis(final_answer="f")],
            ),
            "anthropic": FakeProvider("anthropic", ["b"]),
        },
        default_registry(),
        flagship_models={"openai": "f", "anthropic": "f"},
        cheap_models={"openai": "f", "anthropic": "f"},
    )


def _upload_master(client):
    doc = DocxDocument()
    doc.add_paragraph(
        "Built Terraform modules for Azure, supported AKS workloads and ran "
        "Azure DevOps pipelines with Jenkins."
    )
    buf = io.BytesIO()
    doc.save(buf)
    return client.post(
        "/documents",
        files={"file": ("master.docx", buf.getvalue(), "application/octet-stream")},
        data={"authority": "master_resume", "title": "Master resume"},
    ).json()


async def test_the_whole_one_step_path(db, monkeypatch, tmp_path):
    """Upload a career source, paste a JD, add one line of English, get a DOCX."""
    client = _client(monkeypatch)
    try:
        _upload_master(client)

        r = client.post(
            "/artifacts/resume",
            json={
                "jd_text": GCP_JD,
                "instruction": (
                    "I also have professional Harness experience. Emphasise AKS and "
                    "production troubleshooting and keep it to 2 pages."
                ),
                "name": "Test User",
                "contact": "Charlotte, NC",
            },
        )
        assert r.status_code == 200, r.text
        body = r.json()

        # The outcome decided the workflow; no mode was ever chosen.
        assert body["outcome_kind"] == outcomes.RESUME_TAILOR
        assert body["role_family"]
        assert body["download_url"].endswith("/download")

        # The JD's technologies are reported as gaps, never claimed.
        assert "gcp" in body["gaps"]
        assert "gke" in body["gaps"]

        # And the DOCX actually downloads, and is a real document.
        d = client.get(body["download_url"])
        assert d.status_code == 200
        assert d.headers["content-type"].startswith(
            "application/vnd.openxmlformats-officedocument"
        )
        assert len(d.content) > 5000
        text = "\n".join(
            p.text for p in DocxDocument(io.BytesIO(d.content)).paragraphs
        )
        assert "TEST USER" in text  # the header is upper-cased by the renderer
        assert "Terraform" in text
        # Nothing from the JD that the career does not establish.
        for gap in ("GCP", "GKE", "Cloud Run", "Cloud SQL"):
            assert gap not in text, gap
        # The permanent style rule survives the whole pipeline.
        assert ", and" not in text
    finally:
        client.__exit__(None, None, None)


async def test_a_career_statement_persists_but_a_preference_never_does(
    db, monkeypatch
):
    """The §3 distinction, end to end. 'I have Harness experience' becomes a
    career source; 'keep it to 2 pages' must not become a career fact."""
    client = _client(monkeypatch)
    try:
        _upload_master(client)
        body = client.post(
            "/artifacts/resume",
            json={
                "jd_text": GCP_JD,
                "instruction": (
                    "I also have professional Harness experience. Keep it to 2 pages."
                ),
            },
        ).json()

        assert body["instruction"]["career_statements"] == [
            "I also have professional Harness experience."
        ]
        assert body["instruction"]["preferences"] == ["Keep it to 2 pages."]

        async with session_scope() as s:
            rows = (
                await s.execute(
                    select(DocumentRow).where(
                        DocumentRow.authority == AUTHORITY_USER_STATEMENT
                    )
                )
            ).scalars().all()
        assert len(rows) == 1
        assert "Harness" in rows[0].text
        # The preference is nowhere in stored career evidence.
        assert "2 pages" not in rows[0].text

        # And the statement now genuinely establishes experience.
        confirmed = client.get("/career-profile").json()
        assert "harness" in confirmed["confirmed"]
        assert any(
            s.startswith("user_statement:") for s in confirmed["sources"]["harness"]
        )
    finally:
        client.__exit__(None, None, None)


async def test_the_artifact_is_recorded_with_its_outcome_kind(db, monkeypatch):
    client = _client(monkeypatch)
    try:
        _upload_master(client)
        body = client.post("/artifacts/resume", json={"jd_text": GCP_JD}).json()
        async with session_scope() as s:
            row = await s.get(ArtifactRow, body["id"])
        assert row.kind == outcomes.RESUME_TAILOR
        assert row.cost_usd >= 0
        assert row.trace["instruction"] is not None

        listed = client.get("/artifacts").json()["items"]
        assert listed[0]["id"] == body["id"]
    finally:
        client.__exit__(None, None, None)


async def test_the_stored_file_path_is_never_handed_to_the_client(db, monkeypatch):
    """A server-side path is an implementation detail, and leaking it tells a
    browser where personal documents live on disk."""
    client = _client(monkeypatch)
    try:
        _upload_master(client)
        body = client.post("/artifacts/resume", json={"jd_text": GCP_JD}).json()
        assert "file_path" not in body
        detail = client.get(f"/artifacts/{body['id']}").json()
        assert "file_path" not in detail
        assert detail["download_url"] == f"/artifacts/{body['id']}/download"
    finally:
        client.__exit__(None, None, None)


async def test_a_missing_job_description_is_refused(db, monkeypatch):
    client = _client(monkeypatch)
    try:
        assert client.post("/artifacts/resume", json={"jd_text": "  "}).status_code == 422
        assert (
            client.post(
                "/artifacts/resume", json={"jd_document_id": "nope"}
            ).status_code
            == 404
        )
    finally:
        client.__exit__(None, None, None)


async def test_a_missing_artifact_download_is_reported_not_crashed(db, monkeypatch):
    client = _client(monkeypatch)
    try:
        assert client.get("/artifacts/nope/download").status_code == 404
    finally:
        client.__exit__(None, None, None)


# --------------------------------------------------- instruction parsing


def test_ambiguity_defaults_to_preference_not_to_a_career_claim():
    """A misfiled preference is forgotten after this run. A misfiled career
    fact becomes a permanent claim the user never made. The asymmetry decides
    the default."""
    parsed = parse("Make it sound more senior.")
    assert parsed.career_statements == []
    assert parsed.preferences == ["Make it sound more senior."]


def test_a_first_person_request_is_not_a_career_claim():
    parsed = parse("I want you to emphasise Terraform.")
    assert parsed.career_statements == []


def test_an_empty_instruction_is_harmless():
    assert parse(None).as_dict() == {
        "career_statements": [],
        "preferences": [],
        "denials": [],
    }
    assert parse("   ").preferences == []
    assert parse("   ").denials == []


def test_preferences_reach_the_prompt_labelled_as_preferences():
    """A model told 'target SRE' with no label may read it as a career fact."""
    workflow = _fake_workflow()
    workflow._instruction = parse("Target this as an SRE role.")
    context = workflow._with_preferences("CONFIRMED CAREER CONTEXT: terraform")
    assert "REQUEST-ONLY PREFERENCES" in context
    assert "NOT career facts" in context
    assert "SRE" in context


def test_no_preferences_leaves_the_context_untouched():
    workflow = _fake_workflow()
    workflow._instruction = parse("")
    assert workflow._with_preferences("X") == "X"


# ------------------------------------------- degradation, not destruction


async def test_a_discovery_outage_degrades_instead_of_failing_the_resume(db, monkeypatch):
    """Found by a real run: a provider outage in the OPTIONAL discovery stage
    returned HTTP 500 and produced no resume at all.

    Discovery only widens gap reporting to terms the local vocabulary has never
    seen. The resume is fully writable without it, so an outage must degrade.
    Failing open is safe in the one direction that matters: a missing discovery
    can only under-report a gap, never manufacture a claim, because only career
    evidence confirms anything.
    """
    from council.providers.base import ProviderError

    class DeadProvider(FakeProvider):
        async def generate(self, *a, **k):
            raise ProviderError("Connection error.")

    def workflow_with_dead_discovery():
        workflow = _fake_workflow()
        workflow.providers["reviewer"] = DeadProvider("reviewer")
        return workflow

    import council.engine.factory as factory

    monkeypatch.setattr(factory, "build_resume_workflow", workflow_with_dead_discovery)
    client = TestClient(app, raise_server_exceptions=True)
    client.__enter__()
    try:
        _upload_master(client)
        r = client.post(
            "/artifacts/resume",
            json={"jd_text": GCP_JD + " Also requires Skyforge and Nimbusdeck."},
        )
        assert r.status_code == 200, r.text
        body = r.json()
        # The resume exists and downloads.
        assert client.get(body["download_url"]).status_code == 200
        # Mechanical classification still worked, so known gaps are still known.
        assert "gcp" in body["gaps"]
    finally:
        client.__exit__(None, None, None)


def test_an_unavailable_discovery_is_reported_not_silently_empty():
    from council.documents.discovery import DiscoveryResult

    result = DiscoveryResult(unavailable=True)
    assert result.as_dict()["discovery_unavailable"] is True
    assert result.escalated is False
